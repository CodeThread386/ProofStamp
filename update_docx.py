import docx

doc = docx.Document('ProofStamp_Process_Flows.docx')

replacements = [
    ("User clicks Sign Up (Google OAuth)", "User clicks Sign Up (Google OAuth) and completes Aadhaar eKYC via DigiLocker / Setu API for a legally binding identity."),
    ("Proof Passport Created  —  System creates the user's Proof Passport record in the database:", "Proof Passport Created  —  System creates the user's Proof Passport record in the database (Identity verified via Aadhaar eKYC, ensuring rigorous BSA 2023 Section 63 compliance):"),
    ("Track A — Metadata Injection", "Track A — Metadata Injection (C2PA)"),
    ("System injects a structured metadata block into the file's existing metadata layer", "System injects a structured metadata block using C2PA media provenance standards into the file's existing metadata layer"),
    ("ProofStamp Certificate — a downloadable PDF certificate of ownership", "ProofStamp Certificate — a downloadable PDF certificate, complying with BSA Section 63 Part B, featuring an Aadhaar-based Expert Signature applied via Setu eSign API."),
    ("Database Record Created  —  A stamp record is saved to the database:", "Database Record Created  —  A stamp record is saved to the database (Data Retention and Key Archival SLA policies enforced):"),
    ("Authentication", "Authentication & Legal Infrastructure"),
    ("JWT tokens for session management", "JWT tokens for session management\nSetu eSign API and DigiLocker for Aadhaar eKYC and Expert Signatures"),
    ("crypto module for RSA signing and verification", "crypto module for RSA signing and verification\nAutomated legal notice systems and robust circuit-breaker patterns for reliable eKYC")
]

for para in doc.paragraphs:
    for old_text, new_text in replacements:
        if old_text in para.text:
            # simple replacement
            para.text = para.text.replace(old_text, new_text)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for old_text, new_text in replacements:
                    if old_text in para.text:
                        para.text = para.text.replace(old_text, new_text)

doc.save('ProofStamp_Process_Flows.docx')
